#!/usr/bin/env python3
"""Build a CT Monthly Wrap-Up .docx from a JSON content spec.

Embeds the logo (assets/image1.png) and header banner (assets/image2.png) at the
top, exactly like the example editions, then lays out the title, opener, insight
bullets (bold lead-in + body + hyperlinked source), the Actionable Insight, and
the closing boilerplate.

Upload the resulting .docx to Google Drive as a Google Doc (Drive converts it and
KEEPS the embedded images) -> this is how we get images into the final doc, since
Drive's HTML import cannot embed local files.

Usage:
    python3 build_wrapup_docx.py content.json out.docx [assets_dir]

Content JSON shape:
{
  "title": "Some of Our Most Important Ideas From June, In 2 Minutes",
  "opener": "You survived another month ...",     # the joke; may sound confident
  "transition": "A few ideas worth keeping:",
  "bullets": [
    {"lead": "Bold lead-in claim.",
     "body": "2-4 sentences. Use **bold** for emphasis if needed.",
     "label": "📝 Article: ",             # emoji + type + ': '
     "source_title": "The Strangest Things that Correlate with IQ",
     "source_url": "https://..."},
    ...
  ],
  "actionable_title": "June's Actionable Insight: Write a plan for a recurring problem",
  "actionable_body": "Steps with **bold** labels ...",
  "actionable_label": "From 📝 Article: ",
  "actionable_source_title": "...",
  "actionable_source_url": "...",
  "closing": ["Thank you ...", "If someone shared ..."]
}

Style rules enforced by convention (author the JSON this way):
- NO em dashes or en dashes anywhere. Use commas, colons, parentheses, or split
  the sentence. This script does not insert any dash characters.
- Hedge confidence on interpretive claims ("might", "tends to", "may", "often"),
  except in the opener (the joke) and for firmly established / research-backed facts.
- Each bullet body should rely only on the content of that bullet's own source.
"""
import json
import sys
import os
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CT_BLUE = RGBColor(0x1E, 0x90, 0xFF)


def add_hyperlink(paragraph, url, text):
    """Add a real clickable hyperlink run (blue, underlined) to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1E90FF")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_runs_with_bold(paragraph, text, base_bold=False):
    """Add text to a paragraph, honoring **bold** markup."""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if chunk == "":
            continue
        run = paragraph.add_run(chunk)
        run.bold = base_bold or (i % 2 == 1)


LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")


def add_rich(paragraph, text):
    """Add text honoring inline markdown links [text](url) and **bold**."""
    pos = 0
    for m in LINK_RE.finditer(text):
        add_runs_with_bold(paragraph, text[pos:m.start()])
        add_hyperlink(paragraph, m.group(2), m.group(1))
        pos = m.end()
    add_runs_with_bold(paragraph, text[pos:])


def guard_no_dash(spec):
    """Fail loudly if any em/en dash slipped into the copy."""
    blob = json.dumps(spec, ensure_ascii=False)
    for bad in ("—", "–"):
        if bad in blob:
            raise SystemExit(
                f"Refusing to build: found a '{bad}' (em/en dash) in the content. "
                "Rewrite with commas, colons, parentheses, or a sentence split."
            )


def build(spec, out_path, assets_dir):
    guard_no_dash(spec)
    doc = Document()

    logo = os.path.join(assets_dir, "image1.png")
    header = next(
        (os.path.join(assets_dir, f) for f in ("image2.jpg", "image2.png")
         if os.path.exists(os.path.join(assets_dir, f))),
        os.path.join(assets_dir, "image2.jpg"),
    )

    # Order (baked in): logo wordmark, then the title, then the header banner.
    # Logo (wordmark), centered
    if os.path.exists(logo):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo, width=Inches(2.6))

    # Title, centered, sits between the logo and the header image
    h = doc.add_heading(spec["title"], level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header banner, full content width
    if os.path.exists(header):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(header, width=Inches(6.5))

    # Opener
    doc.add_paragraph(spec["opener"])

    # Transition
    if spec.get("transition"):
        doc.add_paragraph(spec["transition"])

    # Bullets: bold lead-in, then body, then hyperlinked source in parens
    for b in spec["bullets"]:
        p = doc.add_paragraph()
        lead = b["lead"].rstrip()
        if not lead.endswith((".", "?", "!", ":")):
            lead += "."
        p.add_run(lead + " ").bold = True
        add_runs_with_bold(p, b["body"].rstrip() + " ")
        p.add_run("(")
        p.add_run(b["label"])
        add_hyperlink(p, b["source_url"], b["source_title"])
        p.add_run(")")

    # Actionable
    doc.add_heading(spec["actionable_title"], level=1)
    ap = doc.add_paragraph()
    add_rich(ap, spec["actionable_body"])
    src = doc.add_paragraph()
    src.add_run(spec["actionable_label"])
    add_hyperlink(src, spec["actionable_source_url"], spec["actionable_source_title"])

    # Closing boilerplate (honors inline [text](url) links, e.g. Clearer Thinking Plus)
    doc.add_paragraph("")
    for para in spec["closing"]:
        p = doc.add_paragraph()
        add_rich(p, para)

    doc.save(out_path)
    print("wrote", out_path)


if __name__ == "__main__":
    content_path = sys.argv[1]
    out_path = sys.argv[2]
    assets_dir = sys.argv[3] if len(sys.argv) > 3 else os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "assets"
    )
    with open(content_path, encoding="utf-8") as f:
        spec = json.load(f)
    build(spec, out_path, assets_dir)
